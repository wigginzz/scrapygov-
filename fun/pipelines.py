# -*- coding: utf-8 -*-

# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: http://doc.scrapy.org/en/latest/topics/item-pipeline.html
import requests
import docx
from docx import Document
from docx.shared import Inches,RGBColor
from fun import settings
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE
import os
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

'''
    处理，下载
'''
class ImageDownloadPipeline(object):
    def __init__(self): 
	self.file = Document() 
    def process_item(self, item, spider):	
	p = self.file.add_heading(item['title'],level=1)
	paragraph_format = p.paragraph_format
	paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	run = self.file.add_paragraph().add_run(item['url'])
	run.style.name = "Hyperlink"

	if 'image_urls' in item:
		for i in item['image_urls']:
	        	with open('./img.jpg', 'wb') as handle:
	        		response = requests.get(i)
	          		handle.write(response.content)
			self.file.add_picture('./img.jpg',width=Inches(6.0))
			os.remove('./img.jpg')
	if 'image_content' in item:
		for i in item['image_content']:
			p = self.file.add_paragraph()
			str = i.encode('utf-8');
			str = str.replace(' ', '').replace('\n', '').replace('\t', '').replace('\r', '')
			print(len(str))
			if len(str) > 0:
				print(str)
				run = p.add_run(str.decode('utf-8'))
				print('$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$')
				color = RGBColor(0xff, 0x00, 0x00)
				run.font.color.rgb = color
				paragraph_format = p.paragraph_format
				paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		
	for i in item['content']:
		p = self.file.add_paragraph(i)
	return item

    def close_spider(self, spider):
	self.file.save("./word.docx")
